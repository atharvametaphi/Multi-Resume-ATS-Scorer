from __future__ import annotations

from io import BytesIO
from pathlib import Path

import fitz
from docx import Document

from app.core.exceptions import InvalidFileError
from app.utils.file_utils import read_plain_text_file
from app.utils.text_utils import normalize_text


class ExtractionService:
    """Extract raw text from resume and JD files."""

    @staticmethod
    def extract_text(file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            text = ExtractionService._extract_from_pdf(file_path)
        elif suffix == ".docx":
            text = ExtractionService._extract_from_docx(file_path)
        elif suffix == ".txt":
            text = read_plain_text_file(file_path)
        else:
            raise InvalidFileError(f"Unsupported extraction type: {suffix}")
        return normalize_text(text)

    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            text = ExtractionService._extract_from_pdf_bytes(file_bytes)
        elif suffix == ".docx":
            text = ExtractionService._extract_from_docx_bytes(file_bytes)
        elif suffix == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise InvalidFileError(f"Unsupported extraction type: {suffix or 'unknown'}")
        return normalize_text(text)

    @staticmethod
    def _extract_from_pdf(file_path: Path) -> str:
        pages: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                pages.append(page.get_text("text"))
        return "\n".join(pages)

    @staticmethod
    def _extract_from_pdf_bytes(file_bytes: bytes) -> str:
        pages: list[str] = []
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                pages.append(page.get_text("text"))
        return "\n".join(pages)

    @staticmethod
    def _extract_from_docx(file_path: Path) -> str:
        doc = Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)

    @staticmethod
    def _extract_from_docx_bytes(file_bytes: bytes) -> str:
        doc = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)
